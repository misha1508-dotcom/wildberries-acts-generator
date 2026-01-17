#!/usr/bin/env python3
"""
Генератор актов Wildberries
Десктопное приложение с GUI (без веб-сервера)
"""

import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt


class WildberriesActsGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор актов Wildberries")
        self.root.geometry("800x600")
        self.root.resizable(True, True)

        # Данные
        self.current_file = None
        self.months_data = {}

        # Создаем директории
        self.uploads_dir = Path("uploads")
        self.generated_dir = Path("generated")
        self.uploads_dir.mkdir(exist_ok=True)
        self.generated_dir.mkdir(exist_ok=True)

        # Настройка стилей
        self.setup_styles()

        # Создаем интерфейс
        self.create_widgets()

    def setup_styles(self):
        """Настройка стилей для виджетов"""
        style = ttk.Style()
        style.theme_use('default')

        # Стиль для кнопок
        style.configure('Main.TButton',
                       font=('Arial', 11),
                       padding=10)

        style.configure('Action.TButton',
                       font=('Arial', 10, 'bold'),
                       padding=8)

        # Стиль для фреймов
        style.configure('Card.TFrame',
                       background='#f8f9ff',
                       relief='raised')

    def create_widgets(self):
        """Создание элементов интерфейса"""

        # Заголовок
        header_frame = tk.Frame(self.root, bg='#667eea', height=80)
        header_frame.pack(fill='x', pady=(0, 20))
        header_frame.pack_propagate(False)

        title_label = tk.Label(
            header_frame,
            text="Генератор актов Wildberries",
            font=('Arial', 20, 'bold'),
            bg='#667eea',
            fg='white'
        )
        title_label.pack(pady=15)

        subtitle_label = tk.Label(
            header_frame,
            text="Автоматическая генерация актов из XLSX файлов",
            font=('Arial', 10),
            bg='#667eea',
            fg='white'
        )
        subtitle_label.pack()

        # Основной контейнер
        main_container = tk.Frame(self.root, bg='white')
        main_container.pack(fill='both', expand=True, padx=20, pady=10)

        # Секция загрузки файла
        upload_frame = tk.LabelFrame(
            main_container,
            text=" Шаг 1: Загрузите XLSX файл ",
            font=('Arial', 12, 'bold'),
            bg='white',
            fg='#333',
            padx=20,
            pady=20
        )
        upload_frame.pack(fill='x', pady=(0, 20))

        self.file_label = tk.Label(
            upload_frame,
            text="Файл не выбран",
            font=('Arial', 10),
            bg='white',
            fg='#666'
        )
        self.file_label.pack(pady=(0, 10))

        btn_select_file = ttk.Button(
            upload_frame,
            text="📁 Выбрать XLSX файл",
            command=self.select_file,
            style='Main.TButton'
        )
        btn_select_file.pack()

        # Секция выбора месяца
        self.month_frame = tk.LabelFrame(
            main_container,
            text=" Шаг 2: Выберите месяц ",
            font=('Arial', 12, 'bold'),
            bg='white',
            fg='#333',
            padx=20,
            pady=20
        )
        self.month_frame.pack(fill='both', expand=True)

        # Скроллируемая область для месяцев
        canvas = tk.Canvas(self.month_frame, bg='white')
        scrollbar = ttk.Scrollbar(self.month_frame, orient='vertical', command=canvas.yview)
        self.scrollable_frame = tk.Frame(canvas, bg='white')

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        # Изначально показываем подсказку
        self.show_placeholder()

    def show_placeholder(self):
        """Показывает подсказку до загрузки файла"""
        placeholder = tk.Label(
            self.scrollable_frame,
            text="Сначала загрузите XLSX файл",
            font=('Arial', 12),
            bg='white',
            fg='#999'
        )
        placeholder.pack(pady=50)

    def select_file(self):
        """Открывает диалог выбора файла"""
        filename = filedialog.askopenfilename(
            title="Выберите XLSX файл",
            filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")]
        )

        if filename:
            self.current_file = filename
            self.file_label.config(text=f"Файл: {Path(filename).name}")
            self.load_file(filename)

    def load_file(self, filepath):
        """Загружает и парсит XLSX файл"""
        try:
            # Показываем индикатор загрузки
            self.file_label.config(text="Загрузка...")
            self.root.update()

            # Парсим файл
            self.months_data = self.parse_xlsx_file(filepath)

            if not self.months_data:
                messagebox.showwarning(
                    "Предупреждение",
                    "В файле не найдено вкладок с данными"
                )
                self.file_label.config(text="Файл не выбран")
                return

            # Обновляем интерфейс
            self.file_label.config(text=f"Загружено: {Path(filepath).name}")
            self.display_months()

        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Не удалось загрузить файл:\n{str(e)}"
            )
            self.file_label.config(text="Ошибка загрузки")

    def parse_xlsx_file(self, filepath):
        """Парсит XLSX файл и извлекает данные по месяцам"""
        wb = load_workbook(filepath, data_only=True)
        months_data = {}

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            # Извлекаем данные из листа
            data = []
            headers = []

            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i == 0:
                    headers = list(row)
                else:
                    if any(cell is not None for cell in row):
                        row_dict = {}
                        for j, cell in enumerate(row):
                            if j < len(headers):
                                row_dict[headers[j] if headers[j] else f'Column_{j}'] = cell
                        data.append(row_dict)

            if data:  # Добавляем только если есть данные
                months_data[sheet_name] = {
                    'headers': headers,
                    'data': data
                }

        wb.close()
        return months_data

    def display_months(self):
        """Отображает список месяцев"""
        # Очищаем предыдущее содержимое
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()

        # Создаем карточки для каждого месяца
        for month_name in self.months_data.keys():
            self.create_month_card(month_name)

    def create_month_card(self, month_name):
        """Создает карточку месяца"""
        # Карточка
        card = tk.Frame(
            self.scrollable_frame,
            bg='#f8f9ff',
            relief='raised',
            borderwidth=1
        )
        card.pack(fill='x', padx=10, pady=5)

        # Внутренний контейнер
        inner = tk.Frame(card, bg='#f8f9ff')
        inner.pack(fill='x', padx=15, pady=15)

        # Название месяца
        month_label = tk.Label(
            inner,
            text=month_name,
            font=('Arial', 14, 'bold'),
            bg='#f8f9ff',
            fg='#333'
        )
        month_label.pack(anchor='w', pady=(0, 10))

        # Информация о количестве записей
        num_records = len(self.months_data[month_name]['data'])
        info_label = tk.Label(
            inner,
            text=f"Записей: {num_records}",
            font=('Arial', 9),
            bg='#f8f9ff',
            fg='#666'
        )
        info_label.pack(anchor='w', pady=(0, 10))

        # Кнопки
        buttons_frame = tk.Frame(inner, bg='#f8f9ff')
        buttons_frame.pack(fill='x')

        btn_acceptance = tk.Button(
            buttons_frame,
            text="📄 Акт приема-передачи",
            command=lambda: self.generate_document('acceptance', month_name),
            font=('Arial', 10),
            bg='#667eea',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            cursor='hand2'
        )
        btn_acceptance.pack(side='left', padx=(0, 10))

        btn_services = tk.Button(
            buttons_frame,
            text="📋 Акт услуг",
            command=lambda: self.generate_document('services', month_name),
            font=('Arial', 10),
            bg='#f5576c',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            cursor='hand2'
        )
        btn_services.pack(side='left')

    def generate_document(self, doc_type, month_name):
        """Генерирует документ"""
        try:
            # Определяем тип документа
            if doc_type == 'acceptance':
                doc_name = "Акт приема-передачи"
                filename = f'Акт_приема_передачи_{month_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            else:
                doc_name = "Акт услуг"
                filename = f'Акт_услуг_{month_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'

            output_path = self.generated_dir / filename

            # Генерируем документ
            month_data = self.months_data[month_name]

            if doc_type == 'acceptance':
                self.generate_acceptance_act(month_name, month_data, output_path)
            else:
                self.generate_services_act(month_name, month_data, output_path)

            # Показываем сообщение об успехе
            result = messagebox.askyesno(
                "Успех!",
                f"{doc_name} создан!\n\n"
                f"Файл: {filename}\n"
                f"Папка: {self.generated_dir.absolute()}\n\n"
                f"Открыть папку с документом?"
            )

            if result:
                # Открываем папку с файлом
                if os.name == 'nt':  # Windows
                    os.startfile(self.generated_dir)
                elif os.name == 'posix':  # macOS/Linux
                    os.system(f'open "{self.generated_dir}"')

        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Не удалось создать документ:\n{str(e)}"
            )

    def generate_acceptance_act(self, month_name, month_data, output_path):
        """Генерирует акт приема-передачи товара"""
        doc = Document()

        # Заголовок
        title = doc.add_heading('АКТ ПРИЕМА-ПЕРЕДАЧИ ТОВАРА', 0)
        title.alignment = 1

        # Информация
        doc.add_paragraph(f'Месяц: {month_name}')
        doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')
        doc.add_paragraph()
        doc.add_paragraph('Настоящий акт составлен о том, что следующие товары были переданы:')
        doc.add_paragraph()

        # Таблица
        if month_data['data']:
            num_cols = len(month_data['headers'])
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Light Grid Accent 1'

            # Заголовки
            header_cells = table.rows[0].cells
            for i, header in enumerate(month_data['headers']):
                header_cells[i].text = str(header) if header else f'Колонка {i+1}'
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Данные
            for row_data in month_data['data'][:100]:
                row_cells = table.add_row().cells
                for i, header in enumerate(month_data['headers']):
                    key = header if header else f'Column_{i}'
                    value = row_data.get(key, '')
                    row_cells[i].text = str(value) if value is not None else ''

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('Подпись передающей стороны')
        doc.add_paragraph()
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('Подпись принимающей стороны')

        doc.save(output_path)

    def generate_services_act(self, month_name, month_data, output_path):
        """Генерирует акт выполненных услуг"""
        doc = Document()

        # Заголовок
        title = doc.add_heading('АКТ ВЫПОЛНЕННЫХ УСЛУГ', 0)
        title.alignment = 1

        # Информация
        doc.add_paragraph(f'Месяц: {month_name}')
        doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')
        doc.add_paragraph()
        doc.add_paragraph('Настоящий акт составлен о том, что следующие услуги были выполнены по реализации товаров на платформе Wildberries:')
        doc.add_paragraph()

        # Статистика
        total_items = len(month_data['data'])
        doc.add_paragraph(f'Общее количество позиций: {total_items}')

        # Суммы
        price_columns = []
        for header in month_data['headers']:
            if header and any(keyword in str(header).lower() for keyword in ['цена', 'сумма', 'стоимость', 'price', 'sum']):
                price_columns.append(header)

        if price_columns and month_data['data']:
            doc.add_paragraph()
            doc.add_paragraph('Детализация по суммам:')

            for col in price_columns[:3]:
                total = 0
                for row in month_data['data']:
                    value = row.get(col, 0)
                    if value and isinstance(value, (int, float)):
                        total += value
                doc.add_paragraph(f'{col}: {total:,.2f} руб.')

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('Услуги выполнены в полном объеме и в установленные сроки.')
        doc.add_paragraph('Заказчик претензий по объему и качеству выполненных услуг не имеет.')
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('Исполнитель (подпись)')
        doc.add_paragraph()
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('Заказчик (подпись)')

        doc.save(output_path)


def main():
    root = tk.Tk()
    app = WildberriesActsGenerator(root)
    root.mainloop()


if __name__ == "__main__":
    main()
