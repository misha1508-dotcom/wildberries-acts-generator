import os
import json
from flask import Flask, render_template, request, send_file, jsonify, session
from werkzeug.utils import secure_filename
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime

app = Flask(__name__)
app.secret_key = 'wildberries_secret_key_change_in_production'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['GENERATED_FOLDER'] = 'generated'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Создаем необходимые папки
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'xlsx'


def parse_xlsx_file(filepath):
    """Парсит XLSX файл и извлекает данные по месяцам"""
    wb = load_workbook(filepath, data_only=True)
    months_data = {}

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]

        # Извлекаем данные из листа
        data = []
        headers = []

        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i == 0:
                headers = list(row)
            else:
                if any(cell is not None for cell in row):
                    row_dict = {}
                    for j, cell in enumerate(row):
                        if j < len(headers):
                            row_dict[headers[j] if headers[j] else f'Column_{j}'] = cell
                    data.append(row_dict)

        months_data[sheet_name] = {
            'headers': headers,
            'data': data
        }

    wb.close()
    return months_data


def generate_acceptance_act(month_name, month_data, output_path):
    """Генерирует акт приема-передачи товара"""
    doc = Document()

    # Заголовок
    title = doc.add_heading('АКТ ПРИЕМА-ПЕРЕДАЧИ ТОВАРА', 0)
    title.alignment = 1  # Центрирование

    # Информация о месяце
    doc.add_paragraph(f'Месяц: {month_name}')
    doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph()

    # Вводная часть
    doc.add_paragraph('Настоящий акт составлен о том, что следующие товары были переданы:')
    doc.add_paragraph()

    # Таблица с данными
    if month_data['data']:
        # Создаем таблицу
        num_cols = len(month_data['headers'])
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Заголовки таблицы
        header_cells = table.rows[0].cells
        for i, header in enumerate(month_data['headers']):
            header_cells[i].text = str(header) if header else f'Колонка {i+1}'
            # Жирный шрифт для заголовков
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Данные таблицы (первые 100 строк для примера)
        for row_data in month_data['data'][:100]:
            row_cells = table.add_row().cells
            for i, header in enumerate(month_data['headers']):
                key = header if header else f'Column_{i}'
                value = row_data.get(key, '')
                row_cells[i].text = str(value) if value is not None else ''
    else:
        doc.add_paragraph('Нет данных для отображения.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Подписи
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('Подпись передающей стороны')
    doc.add_paragraph()
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('Подпись принимающей стороны')

    # Сохраняем документ
    doc.save(output_path)


def generate_services_act(month_name, month_data, output_path):
    """Генерирует акт выполненных услуг"""
    doc = Document()

    # Заголовок
    title = doc.add_heading('АКТ ВЫПОЛНЕННЫХ УСЛУГ', 0)
    title.alignment = 1  # Центрирование

    # Информация о месяце
    doc.add_paragraph(f'Месяц: {month_name}')
    doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph()

    # Вводная часть
    doc.add_paragraph('Настоящий акт составлен о том, что следующие услуги были выполнены по реализации товаров на платформе Wildberries:')
    doc.add_paragraph()

    # Считаем суммарные показатели
    total_items = len(month_data['data'])

    doc.add_paragraph(f'Общее количество позиций: {total_items}')

    # Пытаемся найти колонки с суммами/ценами
    price_columns = []
    for header in month_data['headers']:
        if header and any(keyword in str(header).lower() for keyword in ['цена', 'сумма', 'стоимость', 'price', 'sum']):
            price_columns.append(header)

    if price_columns and month_data['data']:
        doc.add_paragraph()
        doc.add_paragraph('Детализация по суммам:')

        for col in price_columns[:3]:  # Берем первые 3 колонки с суммами
            total = 0
            for row in month_data['data']:
                value = row.get(col, 0)
                if value and isinstance(value, (int, float)):
                    total += value
            doc.add_paragraph(f'{col}: {total:,.2f} руб.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Заключение
    doc.add_paragraph('Услуги выполнены в полном объеме и в установленные сроки.')
    doc.add_paragraph('Заказчик претензий по объему и качеству выполненных услуг не имеет.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Подписи
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('Исполнитель (подпись)')
    doc.add_paragraph()
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('Заказчик (подпись)')

    # Сохраняем документ
    doc.save(output_path)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Файл не выбран'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Разрешены только XLSX файлы'}), 400

    # Сохраняем файл
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    try:
        # Парсим файл
        months_data = parse_xlsx_file(filepath)

        # Сохраняем данные в сессии
        session['uploaded_file'] = filename
        session['months'] = list(months_data.keys())

        # Сохраняем данные в JSON файл для последующего использования
        json_filename = filename.rsplit('.', 1)[0] + '.json'
        json_filepath = os.path.join(app.config['UPLOAD_FOLDER'], json_filename)

        with open(json_filepath, 'w', encoding='utf-8') as f:
            json.dump(months_data, f, ensure_ascii=False, indent=2, default=str)

        return jsonify({
            'success': True,
            'months': list(months_data.keys())
        })

    except Exception as e:
        return jsonify({'error': f'Ошибка при обработке файла: {str(e)}'}), 500


@app.route('/generate/<doc_type>/<month>')
def generate_document(doc_type, month):
    if 'uploaded_file' not in session:
        return jsonify({'error': 'Сначала загрузите файл'}), 400

    # Загружаем данные из JSON
    filename = session['uploaded_file']
    json_filename = filename.rsplit('.', 1)[0] + '.json'
    json_filepath = os.path.join(app.config['UPLOAD_FOLDER'], json_filename)

    try:
        with open(json_filepath, 'r', encoding='utf-8') as f:
            months_data = json.load(f)

        if month not in months_data:
            return jsonify({'error': 'Месяц не найден'}), 404

        # Генерируем документ
        output_filename = f'{doc_type}_{month}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        output_path = os.path.join(app.config['GENERATED_FOLDER'], output_filename)

        if doc_type == 'acceptance':
            generate_acceptance_act(month, months_data[month], output_path)
        elif doc_type == 'services':
            generate_services_act(month, months_data[month], output_path)
        else:
            return jsonify({'error': 'Неизвестный тип документа'}), 400

        return send_file(output_path, as_attachment=True, download_name=output_filename)

    except Exception as e:
        return jsonify({'error': f'Ошибка при генерации документа: {str(e)}'}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
